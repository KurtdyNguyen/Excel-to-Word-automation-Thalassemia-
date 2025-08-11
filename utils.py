import os
from docxtpl import DocxTemplate
from docxtpl import RichText
from docx import Document
from docx.shared import RGBColor, Pt
import re

TEMPLATES = {
    "PGD": {
        "path": "./templates/template_PGD.docx",
        "output_format": "docx"
    },

    "Thalassemia":{
        "path": "./templates/template_Thalassemia.docx",
        "output_format": "doc"
    }
}

def filename_cleanup(s: str) -> str:
    s = re.sub(r'[<>:"/\\|?*]', '', s)
    s = s.replace('α', 'a').replace('Α', 'A')
    s = s.replace('(', '').replace(')', '')
    return s

def extract_mutation_label(text: str) -> str:
    # Pull out mutation name from formatted sentence
    if "đột biến dị hợp tử" in text:
        match = re.search(r"đột biến dị hợp tử (.+?) trên gen", text)
        if match:
            return match.group(1).strip().upper()
    elif "dị hợp tử" in text:
        match = re.search(r"dị hợp tử (.+?) trên gen", text)
        if match:
            return match.group(1).strip().upper()
    return ""

def render_report(template_type, context, output_name, output_dir, embryos = None):    
    config = TEMPLATES.get(template_type)
    if not config:
        raise FileNotFoundError(f"Không tìm thấy template cho {template_type}")
    
    template_path = config["path"]

    doc = DocxTemplate(template_path)
    doc.render(context)
    
    docx_path = os.path.join(output_dir, output_name + ".docx")
    doc.save(docx_path)

    return docx_path

def highlight_mutation_phrases(docx_path, phrases, color=(255, 0, 0), normal_color=(0, 112, 192)):
    doc = Document(docx_path)
    red = RGBColor(*color)
    blue = RGBColor(*normal_color)
    
    def add_styled_run(paragraph, text, font_color, italic=False, superscript=False):
        run = paragraph.add_run(text)
        run.font.color.rgb = font_color
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.bold = True
        run.italic = italic
        run.font.superscript = superscript
        return run
    
    def add_blue_with_italics(paragraph, text):
        parts = re.split(r"(?i)(HBA|HBB)", text)
        for part in parts:
            if not part:
                continue
            if re.fullmatch(r"(?i)HBA|HBB", part):
                add_styled_run(paragraph, part.upper(), blue, italic=True)
            else:
                add_styled_run(paragraph, part, blue, italic=False)

    for para in doc.paragraphs:
        match_found = False
        for phrase in phrases:
            if phrase in para.text:
                match_found = True
                full_text = para.text
                para.clear()
                before, match, after = full_text.partition(phrase)
                #special case, for "SEA" superscript
                if "--SEA/αα" in full_text:
                    
                    if before:
                        add_blue_with_italics(para, before)

                    first_sea_index = match.find("SEA")
                    if first_sea_index != -1:
                        add_styled_run(para, match[:first_sea_index], red)
                        add_styled_run(para, "SEA", red)
                        
                    second_sea_index = match.find("SEA", first_sea_index + 3)
                    if second_sea_index != -1:
                        add_styled_run(para, match[first_sea_index + 3:second_sea_index], red)
                        add_styled_run(para, "SEA", red, superscript=True)
                        add_styled_run(para, match[second_sea_index + 3:], red)
                    else:
                        add_styled_run(para, match[first_sea_index + 3:], red)
                else:
                    if before:
                        add_blue_with_italics(para, before)
                    add_styled_run(para, match, red)

                if after:
                    add_blue_with_italics(para, after)
                break
        
        if not match_found:
            full_text = para.text
            para.clear()
            add_blue_with_italics(para, full_text)
    doc.save(docx_path)

def extract_red_phrase(sentence: str) -> str:
    if "đột biến dị hợp tử" in sentence:
        start = sentence.find("đột biến dị hợp tử")
        end = sentence.find("trên gen", start)
        return sentence[start:end].strip()
    elif "dị hợp tử" in sentence:
        start = sentence.find("dị hợp tử")
        end = sentence.find("trên gen", start)
        return sentence[start:end].strip()
    return ""